from docx import Document
from docx.shared import Mm, Inches, RGBColor, Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from CORE.tables import set_repeat_table_header, set_cell_vertical_alignment, set_vertical_text, set_cell_margins
import re
from docxtpl import DocxTemplate
import ast

class PMIRenderer:
    def __init__(self, template_path):
        self.doc = Document(template_path)


    @property
    def doc(self):
        """Геттер: возвращает текущий документ"""
        return self._doc

    @doc.setter
    def doc(self, value):
        """Сеттер: устанавливает новый документ с простой проверкой"""
        if value is None:
            raise ValueError("Документ не может быть None")
        self._doc = value


    def add_table_matrix(self, matrix, type=1):
        """
        Добавляет таблицу матрицы входов выходов
        
        """
        self.doc.add_paragraph("Назначение дискретных входов и выходных реле", style='ЮИ_Таблица_Название')
        
        # Создаем таблицу
        table = self.doc.add_table(rows=1, cols=3)
        header_row = table.rows[0]
        header_row.height = Mm(5)
        set_repeat_table_header(header_row)

        # Заголовки столбцов
        if type==1:
            headers = ['Назначенный дискретный вход', 'Обозначение сигнала', 'Назначенное выходное реле']
        else:
            headers = ['Назначенный дискретный вход или функциональная клавиша', 'Обозначение сигнала', 'Назначенный светодиод / Назначенное выходное реле']           
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell.paragraphs[0].style = 'Текст таблицы'
            set_cell_vertical_alignment(cell, align="center")

        # Заполнение данными из JSON
        for item in matrix:
            row = table.add_row()
            
            row.cells[0].text = item[0]
            row.cells[1].text = item[1]
            row.cells[2].text = item[2]

            # Выравнивание
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    paragraph.style = 'Текст таблицы'

        # Стиль таблицы
        table.style = 'Стиль3'
        table.allow_autofit = True
        table.autofit = True
        
        # Ширина столбцов
        #widths = [2, 2, 2]
        #for i, width in enumerate(widths):
            #for cell in table.columns[i].cells:
                #cell.width = Inches(width)    


    def save_docx(self, name='Output/ПМИ.docx'):
        self.doc.save(name)
 

    def doc_add_heading_one(self, heading):
        self.doc.add_heading(heading, level=1)

    def doc_add_heading_two(self, heading):
        self.doc.add_heading(heading, level=2)       

    def add_blanc(self, modes_list):
        #target_doc = doc if doc is not None else self.doc
        
        for mode_index, mode in enumerate(modes_list):
            # ... (ваш код без изменений, но заменяем self.doc на target_doc) ...
            mode_number = mode["mode_file"].replace('.xlsx', '').split('_')[-1]
            self.doc.add_heading(f'Режим №{mode_number}', level=2)

            if mode_index == 0:
                for block_name, setting_block in mode['settings'].items():
                    # Если заголовки блоков нужны только в основном документе, оставьте как есть.
                    # Но обычно они тоже должны идти в target_doc
                    self.doc.add_heading(setting_block["BlockFullRusName"], level=3)
                    
                    for func_name, function in setting_block['functions'].items():
                        # Передаем target_doc внутрь, если _add_table_settings тоже использует self.doc
                        # Если _add_table_settings жестко привязан к self.doc, то вариант с local var сложнее.
                        # В вашем текущем коде _add_table_settings использует self.doc напрямую.
                        # Поэтому вариант с заменой self.doc (как у вас) проще, но требует осторожности.
                        self._add_table_settings(function, mode['mode_file'], is_first_mode=True)
            
            else:
                has_changes = self._has_mode_changes(mode)
                
                if not has_changes:
                    paragraph = self.doc.add_paragraph(style='ЮИ_Обычный')
                    run = paragraph.add_run("Уставки и параметры задействованных функций идентичны предыдущему режиму.")
                    run.font.italic = True
                    continue
                
                for block_name, setting_block in mode['settings'].items():
                    has_block_changes = self._has_block_changes(setting_block)
                    
                    if not has_block_changes:
                        continue
                    
                    for func_name, function in setting_block['functions'].items():
                        has_func_changes = self._has_func_changes(function)
                        if not has_func_changes:
                            continue
                        
                        self._add_table_settings(function, mode['mode_file'], is_first_mode=False)

   


    def _has_mode_changes(self, mode):
        """Проверяет, есть ли изменения в режиме"""
        for block_name, setting_block in mode['settings'].items():
            if self._has_block_changes(setting_block):
                return True
        return False

    def _has_block_changes(self, setting_block):
        """Проверяет, есть ли изменения в блоке"""
        for func_name, function in setting_block['functions'].items():
            if self._has_func_changes(function):
                return True
        return False

    def _has_func_changes(self, function):
        """Проверяет, есть ли изменения в функции"""
        for param in function.get('parameters', []):
            if param.get('Color') == '1':
                return True
        return False

    def _parse_options_string(self, text):
        """
        Преобразует строку вида "1 - Управляющее напряжение, 2 - Вольтметровая блокировка"
        в словарь {"1": "Управляющее напряжение", "2": "Вольтметровая блокировка"}
        """
        if not text or not isinstance(text, str):
            return {} , ""
        
        result = {}
        # Разделяем по запятой (учитываем возможные пробелы)
        pairs = re.split(r'\s*,\s*', text.strip())
        
        for pair in pairs:
            # Разделяем по " - " или "-" с возможными пробелами
            match = re.match(r'^\s*(\d+)\s*-\s*(.+?)\s*$', pair)
            if match:
                key = match.group(1)
                value = match.group(2).strip()
                result[key] = value
        
        return result, " / ".join(result.values())


    def _format_values(self, step_val, default, min_val, max_val, current):
        """
        Форматирует значения по шагу: выравнивает количество знаков и меняет точку на запятую.
        """
        # Определяем количество знаков после запятой по шагу
        decimals = len(step_val.split('.')[1]) if '.' in step_val else 0
        
        # Форматируем каждое значение
        default = f"{float(default):.{decimals}f}".replace('.', ',')
        min_val = f"{float(min_val):.{decimals}f}".replace('.', ',')
        max_val = f"{float(max_val):.{decimals}f}".replace('.', ',')
        step = f"{float(step_val):.{decimals}f}".replace('.', ',')
        current = f"{float(current):.{decimals}f}".replace('.', ',')
        
        return step, default, min_val, max_val, current


    def _add_table_settings(self, function, mode_file, is_first_mode=True):
        """Добавляет таблицу настроек"""
        func_name = function.get("NodeNameRu", "")
        func_full_name = function.get("FullNodeNameRu", "")

        # Заголовок функции
        header_func = f'{func_name} ({func_full_name})'
        if func_name == 'Общие уставки':
            header_func = 'Общие уставки'
        
        self.doc.add_paragraph(header_func, style='ЮИ_Таблица_Название')
        
        # Создаем таблицу
        table = self.doc.add_table(rows=1, cols=7)
        header_row = table.rows[0]
        header_row.height = Mm(5)
        set_repeat_table_header(header_row)

        # Заголовки столбцов
        headers = ['Параметр', 'Обозначение ФСУ', 'Значение / Диапазон', 'Ед. изм.', 'Шаг', 'Значение по умолчанию', 'Выставленная уставка']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell.paragraphs[0].style = 'Текст таблицы'
            set_cell_vertical_alignment(cell, align="center")

        # Заполнение данными
        for item in function.get("parameters", []):
            # ✅ Для не-первого режима пропускаем параметры без изменений
            if not is_first_mode and item.get('Color') != '1':
                continue

            # подготовка значений для вставки в таблицу
            note = item.get('Note')
            default = item.get('DefaultValue')
            current = str(item.get('CurrentSetting'))
            step = item.get('Step')

            diap = ""
            if note != '':
                option_dict, diap = self._parse_options_string(note)
                default_setting = option_dict.get(default)
                current_setting = option_dict.get(current)
                step = '-'
                
            else:
                #print(func_full_name, func_name, item.get('FullDescription'))
                step_val, default_val, min_val, max_val, current_val = self._format_values(step, default, item.get('Min'), item.get('Max'), current)
                diap = f"{min_val} ... {max_val}"
                default_setting = default_val
                current_setting = current_val
                step = step_val


            row = table.add_row()
            row.cells[0].text = f"{item.get('FullDescription', '')} ({item.get('Description', '')}) "
            row.cells[1].text = item.get('AppliedDescription') if item.get('AppliedDescription')!="" else "-"
            row.cells[2].text = diap
            row.cells[3].text = item.get('Units') if item.get('Units')!="" else "-"
            row.cells[4].text = step
            row.cells[5].text = default_setting
            row.cells[6].text = current_setting

            # Жирный шрифт для ячейки "Уставка"
            cell_ust = row.cells[6]
            for paragraph in cell_ust.paragraphs:
                for run in paragraph.runs:
                    #run.bold = True
                    run.italic = True

            # Выравнивание
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    paragraph.style = 'Текст таблицы'

            # Цветовая индикация (только для изменённых)
            if item.get('Color') == '1':
                shading_color = "00FF7B" #"11502F" #"FFC000"  # Желтый
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shading_color}"/>')
                cell_ust._tc.get_or_add_tcPr().append(shading_elm)

        table.style = 'Стиль3'
        table.allow_autofit = False
        table.autofit = False
        
        # Ширина столбцов
        widths = [1.35, 0.98, 1.35, 0.48, 0.47, 1.0, 1.05]
        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(width)


    def add_inout_table(self, df_ins, type):
        """
        Добавляет таблицу входных воздействий в документ на основе DataFrame.
        """
        from docx.shared import Mm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if type == 'inputs':
            self.doc.add_paragraph("Подаваемые воздействия при проверке", style='ЮИ_Таблица_Название')
        else:
            self.doc.add_paragraph("Контролируемые сигналы при проверке", style='ЮИ_Таблица_Название')

        # 1. Формируем заголовки: № режима + колонки DataFrame
        headers = ["Номер режима"] + list(df_ins.columns)

        # 2. Создаем таблицу
        table = self.doc.add_table(rows=1, cols=len(headers))
        header_row = table.rows[0]
        #set_repeat_table_header(header_row)

        # 3. Заполняем заголовки
        max_len_str = 0
        for i, header_text in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = str(header_text)

            if len(header_text) > max_len_str:
                max_len_str = len(header_text)
            
            # Стилизация заголовка
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].style = 'Текст таблицы'
            set_vertical_text(cell)  # Поворот текста
            set_cell_vertical_alignment(cell, align="center")
            set_cell_margins(cell, top=0.0, bottom=0.0, left=0.0, right=0.0)  # Убираем лишние отступы

        # 4. Заполняем строки данными из DataFrame
        for idx, row in df_ins.iterrows():
            # Основная строка с данными
            new_row = table.add_row()
            
            # Первый столбец — номер режима (индекс + 1)
            new_row.cells[0].text = str(idx + 1)
            
            # Остальные столбцы — значения из строки DataFrame
            for col_idx, value in enumerate(row.values):
                value_str = str(value).replace('.', ',') # Замена . на , в числах таблицы входов
                new_row.cells[col_idx + 1].text = value_str
            
            # Стилизация ячеек основной строки
            for cell in new_row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.style = 'Текст таблицы'
            
            # Добавление пустой строки ТОЛЬКО для type = 'outputs'
            if type == 'outputs':
                empty_row = table.add_row()
                
                # Первый столбец пустой строки — тот же номер режима
                empty_row.cells[0].text = str(idx + 1)
                
                # Остальные ячейки пустые
                for col_idx in range(1, len(empty_row.cells)):
                    empty_row.cells[col_idx].text = ""
                
                # Стилизация пустой строки
                for cell in empty_row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph.style = 'Текст таблицы'

        # 5. Настройка стиля и ширины таблицы
        table.style = 'Стиль3'
        table.allow_autofit = False

        header_row.height = Mm(15 + max_len_str) 
        set_repeat_table_header(header_row)
            # Ширина столбцов
            #widths = [2, 2, 2]
            #for i, width in enumerate(widths):
                #for cell in table.columns[i].cells:
                    #cell.width = Inches(width)      

    def insert_settings_blancs(self, blanc):
        for mode_num, mode_data in blanc.final_output.items():

            if not mode_data or mode_data==[]:
                continue
            self.doc.add_heading(f'Режим №{mode_num}', level=2)

            for data in mode_data:
                # Заголовок функции
                self.doc.add_paragraph(data["func_name"].split("_")[0], style='ЮИ_Таблица_Название')
                
                # Создаем таблицу
                table = self.doc.add_table(rows=1, cols=7)
                header_row = table.rows[0]
                header_row.height = Mm(5)
                set_repeat_table_header(header_row)
                
                # Заголовки столбцов
                headers = ['Параметр', 'Обозначение ФСУ', 'Значение / Диапазон', 'Ед. изм.', 'Шаг', 'Значение по умолчанию', 'Выставленная уставка']
                for i, header in enumerate(headers):
                    cell = table.cell(0, i)
                    cell.text = header
                    cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    cell.paragraphs[0].style = 'Текст таблицы'
                    set_cell_vertical_alignment(cell, align="center")
                
                row_count = 1  # Учитываем строку заголовка
                
                if data["type"] == "simple":
                    # Простая функция
                    for item in data["rows"]:
                        row = table.add_row()


                        if item["col3"].startswith("note_"):
                            third_col = item["col3"].replace("note_", "", 1) # 1 означает заменить только первое вхождение
                            dict_note = ast.literal_eval(third_col)
                            third = ""
                            for value in dict_note.values():
                                third +=value +' /\n'
                            third = third[:-3]

                            seven = dict_note.get(item["col7"])
                        else:
                            third = item["col3"] # Или какое-то значение по умолчанию
                            #seven = item["col7"]

                            step = float(item["col5"].replace(",", "."))
                            step_str = str(step).rstrip('0').rstrip('.')
                            decimals = len(step_str.split('.')[1]) if '.' in step_str else 0
                            seven = f"{step:.{decimals}f}".replace('.', ',')


                        row.cells[0].text = item["col1"]
                        row.cells[1].text = item["col2"] if item["col2"] != "" else "-"
                        row.cells[2].text = third
                        row.cells[3].text = item["col4"] if item["col4"] != "" else "-"
                        row.cells[4].text = item["col5"] if item["col5"] != "" else "-"
                        row.cells[5].text = item["col6"]
                        row.cells[6].text = seven
                        
                        # Форматирование ячеек
                        cell_ust = row.cells[6]
                        for paragraph in cell_ust.paragraphs:
                            for run in paragraph.runs:
                                #run.bold = True
                                run.italic = True
                        
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                paragraph.style = 'Текст таблицы'
                        
                        if mode_num != '1':
                            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="00FF7B"/>')
                            cell_ust._tc.get_or_add_tcPr().append(shading_elm)
                        
                elif data["type"] == "complex":
                    # Сложная функция - добавляем разделители для каждой подфункции
                    for sub_func in data["sub_functions"]:
                        subtitle = sub_func["subtitle"]
                        
                        # Добавляем строку-разделитель с заголовком подфункции
                        row = table.add_row()

                        # Объединяем все 7 ячеек в одну для красивого заголовка
                        # merge возвращает объект Cell, который охватывает весь диапазон
                        merged_cell = row.cells[0].merge(row.cells[6])
                        
                        merged_cell.text = f"▼ {subtitle}"
                        
                        # Настраиваем стиль объединенной ячейки
                        for paragraph in merged_cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.size = Pt(10)  # <--- Устанавливаем размер шрифта 10
                                #run.font.name = 'Arial Narrow' # Опционально, чтобы шрифт был единообразным
                        
                        # Закрашиваем фон объединенной ячейки
                        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>')
                        merged_cell._tc.get_or_add_tcPr().append(shading_elm)
                        
                        # Вертикальное выравнивание для объединенной ячейки
                        set_cell_vertical_alignment(merged_cell, align="center")

                        # Добавляем строки параметров подфункции
                        for item in sub_func["rows"]:

                            if item["col3"].startswith("note_"):
                                third_col = item["col3"].replace("note_", "", 1) # 1 означает заменить только первое вхождение
                                dict_note = ast.literal_eval(third_col)
                                third = ""
                                for value in dict_note.values():
                                    third +=value +' /\n'
                                third = third[:-3]

                                seven = dict_note.get(item["col7"])
                            else:
                                third = item["col3"] # Или какое-то значение по умолчанию
                                #seven = item["col7"]

                                step = float(item["col5"].replace(",", "."))
                                step_str = str(step).rstrip('0').rstrip('.')
                                decimals = len(step_str.split('.')[1]) if '.' in step_str else 0

                                # Преобразуем строку в число перед форматированием
                                col7_value = float(item["col7"].replace(",", "."))
                                seven = f"{col7_value:.{decimals}f}".replace('.', ',')


                            row = table.add_row()
                            row.cells[0].text = item["col1"]
                            row.cells[1].text = item["col2"] if item["col2"] != "" else "-"
                            row.cells[2].text = third 
                            row.cells[3].text = item["col4"] if item["col4"] != "" else "-"
                            row.cells[4].text = item["col5"] if item["col5"] != "" else "-"
                            row.cells[5].text = item["col6"]
                            row.cells[6].text = seven
                            
                            cell_ust = row.cells[6]
                            for paragraph in cell_ust.paragraphs:
                                for run in paragraph.runs:
                                    #run.bold = True
                                    run.italic = True
                            
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    paragraph.style = 'Текст таблицы'
                            
                            if mode_num != '1':
                                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="00FF7B"/>')
                                cell_ust._tc.get_or_add_tcPr().append(shading_elm)
                
                # Настройка таблицы
                table.style = 'Стиль3'
                table.allow_autofit = False
                table.autofit = False
                
                widths = [1.35, 0.98, 1.35, 0.48, 0.47, 1.0, 1.05]
                for i, width in enumerate(widths):
                    for cell in table.columns[i].cells:
                        cell.width = Inches(width)
        