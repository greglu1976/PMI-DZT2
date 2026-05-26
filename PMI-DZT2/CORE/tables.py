# Вынесена генерация таблиц
from docx.shared import Pt, Mm

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

from lxml import etree

# Функция для применения стиля ко всем параграфам в ячейках таблицы
def apply_style_to_all_cells(table, style_name, numbered_style="ЮИ_Таблица_Нумерованный"):
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            # Пропускаем первую строку (заголовок таблицы)
            if row_idx == 0:
                continue  # Не применяем стили к строке заголовка

            for paragraph in cell.paragraphs:
                # Применяем специальный стиль для первого столбца
                if col_idx == 0:  # Проверяем, является ли это ячейкой первого столбца
                    paragraph.style = numbered_style
                else:
                    # Применяем указанный стиль ко всем остальным ячейкам
                    paragraph.style = style_name

def set_cell_vertical_alignment(cell, align="center"):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcValign = OxmlElement('w:vAlign')
        tcValign.set(qn('w:val'), align)
        tcPr.append(tcValign)

# Функция для установки отступов в ячейке
def set_cell_margins(cell, top=0.05, bottom=0.05, left=0.05, right=0.05):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    tcMar = etree.SubElement(tcPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcMar', nsmap=nsmap)
    
    for position, value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        element = etree.SubElement(tcMar, f'{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}{position}')
        element.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', str(int(value * 1440)))  # 1440 EMUs в 1 мм
        element.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'dxa')

# Функция для установки вертикальной ориентации текста
def set_vertical_text(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    text_direction = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}textDirection')
    text_direction.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'btLr')
    tcPr.append(text_direction)

def set_repeat_table_header(row):
    """ set repeat table row on every new page
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row

# Функция для форматирования числовых значений
def format_number(value):
    try:
        # Пробуем преобразовать в число
        num = float(str(value).replace(',', '.'))
        if num.is_integer():
            return str(int(num))  # Без дробной части
        return str(num).replace('.', ',')  # С дробной частью
    except (ValueError, TypeError):
        return str(value)  # Если не число, возвращаем как есть

def add_table_infuences(doc, table_rows):

    doc.add_paragraph('Подаваемые воздействия при проверке', style='ЮИ_Таблица_Название')
    table = doc.add_table(rows=1, cols=len(table_rows[0]))

    ##################################################
    # Рассчитываем высоту заголовка
    header_row_height = 25  # Базовая высота (для 'Номер режима')
    max_length = max(len(str(header)) for header in table_rows[0])
    # Проверяем необходимость увеличения высоты
    if len(table_rows[0]) < 14:  # Если столбцов <= 17, не меняем высоту
        pass
    elif len(table_rows[0]) >= 14 and len(table_rows[0]) <= 17: 
        # Ищем самый длинный заголовок
        if max_length > 22:
            # Увеличиваем высоту на 2мм за каждый символ сверх 12
            additional_height = (max_length - 22) * 2
            header_row_height += additional_height
    else:
        # Ищем самый длинный заголовок
        if max_length > 12:
            # Увеличиваем высоту на 2мм за каждый символ сверх 12
            additional_height = (max_length - 12) * 2
            header_row_height += additional_height

    # Установка высоты первой строки (заголовок)
    header_row = table.rows[0]
    header_row.height = Mm(header_row_height)  # Устанавливаем высоту строки заголовка
    ##################################################

    set_repeat_table_header(header_row)

    # Добавление заголовков
    for i, header in enumerate(table_rows[0]):
        cell = header_row.cells[i]

        paragraph = cell.paragraphs[0]
        paragraph.text = header
        paragraph.style = 'Текст таблицы'  # Применяем стиль
        #paragraph.runs[0].bold = True      # Делаем жирным (если нужно)

        #run = cell.paragraphs[0].add_run(header) # так было
        set_vertical_text(cell)
        set_cell_vertical_alignment(cell, align="center")
        set_cell_margins(cell, top=0.0, bottom=0.0, left=0.0, right=0.0)

    # Добавление данных
    for row_data in table_rows[1:]:
        row_cells = table.add_row().cells
        for j, value in enumerate(row_data):
            cell = row_cells[j]
            cell.text = format_number(value)
            set_cell_margins(cell, top=0.0, bottom=0.0, left=0.0, right=0.0)
    
    # Применяем стиль ко всем ячейкам

    apply_style_to_all_cells(table, 'Текст таблицы', numbered_style='Текст таблицы')
    table.style = 'Стиль3'  # Применение стиля таблицы из шаблона

    table.allow_autofit = False
    table.autofit = False
    table.style = 'Стиль3'  # Применение стиля таблицы из шаблона

    # Узнаем количество столбцов
    num_columns = len(table.columns)
    # Одна ширина для всех столбцов - адаптивная
    #column_width = max(0.2, 5.0 / num_columns)  # Минимум 0.5 дюйма, максимум ~1.7 дюйма
    column_width = 6.674 / num_columns
    # Применяем ширину
    for column in table.columns:
        for cell in column.cells:
            cell.width = Inches(column_width)

    return doc

def add_table_results(doc, table_rows):

    doc.add_paragraph('Контролируемые сигналы при проверке', style='ЮИ_Таблица_Название')
    table = doc.add_table(rows=1, cols=len(table_rows[0]))
    table.style = 'Стиль3'  # Применение стиля таблицы из шаблона

    # Рассчитываем высоту заголовка
    header_row_height = 25  # Базовая высота (для 'Номер режима')
    max_length = max(len(str(header)) for header in table_rows[0])
    # Проверяем необходимость увеличения высоты
    if len(table_rows[0]) < 14:  # Если столбцов <= 17, не меняем высоту
        pass
    elif len(table_rows[0]) >= 14 and len(table_rows[0]) <= 17: 
        # Ищем самый длинный заголовок
        if max_length > 22:
            # Увеличиваем высоту на 2мм за каждый символ сверх 12
            additional_height = (max_length - 22) * 2
            header_row_height += additional_height
    else:
        # Ищем самый длинный заголовок
        if max_length > 12:
            # Увеличиваем высоту на 2мм за каждый символ сверх 12
            additional_height = (max_length - 12) * 2
            header_row_height += additional_height

    # Установка высоты первой строки (заголовок)
    header_row = table.rows[0]
    header_row.height = Mm(header_row_height)  # Устанавливаем высоту строки заголовка
    set_repeat_table_header(header_row)

    # Добавление заголовков
    for i, header in enumerate(table_rows[0]):
        cell = header_row.cells[i]

        paragraph = cell.paragraphs[0]
        paragraph.text = header
        paragraph.style = 'Текст таблицы'  # Применяем стиль
        #paragraph.runs[0].bold = True      # Делаем жирным (если нужно)

        #run = cell.paragraphs[0].add_run(header)
        set_vertical_text(cell)
        set_cell_vertical_alignment(cell, align="center")
        set_cell_margins(cell, top=0.0, bottom=0.0, left=0.0, right=0.0)

    # Добавляем строки с данными и пустые строки после КАЖДОЙ строки
    for row_data in table_rows[1:]:
        # 1. Добавляем строку с данными
        data_row = table.add_row()
        for j, value in enumerate(row_data):
            cell = data_row.cells[j]
            cell.text = format_number(value)
            set_cell_margins(cell, top=0.0, bottom=0.0, left=0.0, right=0.0)
        
        # 2. Добавляем пустую строку с тем же номером режима
        empty_row = table.add_row()
        for j in range(len(table_rows[0])):
            cell = empty_row.cells[j]
            if j == 0:  # Первый столбец - номер режима
                cell.text = str(row_data[0])  # Сохраняем номер режима
            else:
                cell.text = ""  # Пустые ячейки
            set_cell_margins(cell, top=0.0, bottom=0.0, left=0.0, right=0.0)

    # Применяем стиль ко всем ячейкам
    apply_style_to_all_cells(table, 'Текст таблицы', numbered_style='Текст таблицы')

    table.allow_autofit = False
    table.autofit = False
    table.style = 'Стиль3'  # Применение стиля таблицы из шаблона

    # Узнаем количество столбцов
    num_columns = len(table.columns)
    # Одна ширина для всех столбцов - адаптивная
    #column_width = max(0.2, 5.0 / num_columns)  # Минимум 0.5 дюйма, максимум ~1.7 дюйма
    column_width = 6.674 / num_columns
    # Применяем ширину
    for column in table.columns:
        for cell in column.cells:
            cell.width = Inches(column_width)
    return doc


##### ДЛЯ БЛАНКА УСТАВОК ПМИ

def add_table_settingsOLD(doc, data_list, descriptions):
    #print(data_list)
    if not data_list:
        return doc
   
    key1 = data_list[0]['LD'].lower()
    key2 = data_list[0]['LN'].lower()
    #print('>>>', key1, key2)
    func_name = descriptions[key1][key2]['funcname']
    func_short_name = descriptions[key1][key2]['func_short_name']

    #doc.add_heading(key_desc, level=4)
    header_func = func_name + f' ({func_short_name})'
    if func_name == 'Общие уставки':
        header_func = 'Общие уставки'
    
    doc.add_paragraph(header_func, style='ЮИ_Таблица_Название')
    # Создаем таблицу
    table = doc.add_table(rows=1, cols=7)
    # Установка высоты первой строки (заголовок)
    header_row = table.rows[0]
    header_row.height = Mm(5)  # Устанавливаем высоту строки заголовка в 45 мм
    set_repeat_table_header(header_row)

    # Заголовки столбцов
    headers = ['Параметр', 'Обозначение ФСУ', 'Значение / Диапазон', 'Ед. изм.', 'Шаг', 'Значение по умолчанию', 'Уставка']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        # Делаем заголовки жирными
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.paragraphs[0].style = 'Текст таблицы'
        set_cell_vertical_alignment(cell, align="center")

    for item in data_list:
        row = table.add_row()
        is_changed =item.get('IsChanged', 0)
        
        # Заполнение ячеек
        row.cells[0].text = item.get('Параметр', '')
        row.cells[1].text = item.get('Обозначение ФСУ', '-')
        row.cells[2].text = item.get('Значение / Диапазон', '')
        row.cells[3].text = item.get('Ед.изм.', '-')
        row.cells[4].text = item.get('Шаг', '-')
        row.cells[5].text = item.get('Значение по умолчанию', '')

        cell_ust = row.cells[6]
        cell_ust.text = str(item.get('Уставка', ''))

        # делаем жирным
        for paragraph in cell_ust.paragraphs:
            for run in paragraph.runs:
                run.bold = True

        # Выравнивание текста по центру (если нужно)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                paragraph.style = 'Текст таблицы'


        # Цветовая индикация для ячейки "Уставка"
        if is_changed == 1:
            shading_color = "FFC000"  # Желтый
        elif is_changed == 2:
            shading_color = "FF0000"  # Красный
        else:
            shading_color = None  # Нет заливки

        if shading_color:

            #shading_elm = parse_xml(r'<w:shd {} w:fill="FFC000"/>'.format(nsdecls('w')))
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shading_color}"/>')
            cell_ust._tc.get_or_add_tcPr().append(shading_elm)
            # Добавляем заливку только для ячейки с уставкой
            #tc = cell_ust._tc
            #tcPr = tc.get_or_add_tcPr()
            #shading_elm = tcPr._element.xml
            #shading_elm = f'<w:shd w:fill="{shading_color}"/>'
            #tcPr._element.insert_element_before(shading_elm, "w:tcBorders")



    table.style = 'Стиль3'
    table.allow_autofit = False
    table.autofit = False
    # Задаем ширину столбцов (в дюймах)
    widths = [1.75, 0.98, 1.4, 0.48, 0.47, 1.0, 0.6]
    for i, width in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = Inches(width)    
    return doc






##############################################################

#############################################################


#############################################################

###############################################################






def add_table_settings(doc, function, mode_file=None):
    """
    Добавляет таблицу настроек для функции
    
    Args:
        doc: документ Word
        function: словарь функции из JSON
        mode_file: имя файла режима (для отладки)
    """
    if not function or 'parameters' not in function:
        return doc

    func_name = function.get("NodeNameRu", "")
    func_full_name = function.get("FullNodeNameRu", "")

    # Заголовок функции
    header_func = f'{func_name} ({func_full_name})'
    if func_name == 'Общие уставки':
        header_func = 'Общие уставки'
    
    doc.add_paragraph(header_func, style='ЮИ_Таблица_Название')
    
    # Создаем таблицу
    table = doc.add_table(rows=1, cols=7)
    header_row = table.rows[0]
    header_row.height = Mm(5)
    set_repeat_table_header(header_row)

    # Заголовки столбцов
    headers = ['Параметр', 'Обозначение ФСУ', 'Значение / Диапазон', 'Ед. изм.', 'Шаг', 'Значение по умолчанию', 'Уставка']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.paragraphs[0].style = 'Текст таблицы'
        set_cell_vertical_alignment(cell, align="center")

    # Заполнение данными из JSON
    for item in function.get("parameters", []):
        row = table.add_row()
        
        # !!! ИСПРАВЛЕНО: используем правильные имена полей из JSON !!!
        row.cells[0].text = f"{item.get('FullDescription', '')}  ({item.get('Description', '')})"  # Параметр
        row.cells[1].text = item.get('AppliedDescription', '-')           # Обозначение ФСУ
        row.cells[2].text = f"{item.get('Min', '')}-{item.get('Max', '')}"  # Диапазон
        row.cells[3].text = item.get('Units', '-')                  # Ед. изм.
        row.cells[4].text = item.get('Step', '-')                   # Шаг
        row.cells[5].text = item.get('DefaultValue', '')            # По умолчанию
        row.cells[6].text = str(item.get('CurrentSetting', ''))     # Уставка

        # Жирный шрифт для ячейки "Уставка"
        cell_ust = row.cells[6]
        for paragraph in cell_ust.paragraphs:
            for run in paragraph.runs:
                run.bold = True

        # Выравнивание
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                paragraph.style = 'Текст таблицы'

        # !!! ИСПРАВЛЕНО: Color с большой буквы !!!
        is_changed = item.get('Color', '0')
        
        # Цветовая индикация
        if is_changed == '1' or is_changed == 1:
            shading_color = "FFC000"  # Желтый
        elif is_changed == '2' or is_changed == 2:
            shading_color = "FF0000"  # Красный
        else:
            shading_color = None

        if shading_color:
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shading_color}"/>')
            cell_ust._tc.get_or_add_tcPr().append(shading_elm)

    # Стиль таблицы
    table.style = 'Стиль3'
    table.allow_autofit = False
    table.autofit = False
    
    # Ширина столбцов
    widths = [1.75, 0.98, 1.4, 0.48, 0.47, 1.0, 0.6]
    for i, width in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = Inches(width)    
    
    return doc



def add_table_settings_v2(doc, function):
    """
    Добавляет таблицу настроек для функции
    
    Args:
        doc: документ Word
        function: словарь функции из JSON
    """

    print(function)
    return

    if not function:
        return doc

    func_name = function.get("NodeNameRu", "")
    func_full_name = function.get("FullNodeNameRu", "")

    # Заголовок функции
    header_func = f'{func_name} ({func_full_name})'
    if func_name == 'Общие уставки':
        header_func = 'Общие уставки'
    
    doc.add_paragraph(header_func, style='ЮИ_Таблица_Название')
    
    # Создаем таблицу
    table = doc.add_table(rows=1, cols=7)
    header_row = table.rows[0]
    header_row.height = Mm(5)
    set_repeat_table_header(header_row)

    # Заголовки столбцов
    headers = ['Параметр', 'Обозначение ФСУ', 'Значение / Диапазон', 'Ед. изм.', 'Шаг', 'Значение по умолчанию', 'Уставка']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.paragraphs[0].style = 'Текст таблицы'
        set_cell_vertical_alignment(cell, align="center")

    # Заполнение данными из JSON
    for item in function.get("parameters", []):
        row = table.add_row()
        
        # !!! ИСПРАВЛЕНО: используем правильные имена полей из JSON !!!
        row.cells[0].text = f"{item.get('FullDescription', '')}  ({item.get('Description', '')})"  # Параметр
        row.cells[1].text = item.get('AppliedDescription', '-')           # Обозначение ФСУ
        row.cells[2].text = f"{item.get('Min', '')}-{item.get('Max', '')}"  # Диапазон
        row.cells[3].text = item.get('Units', '-')                  # Ед. изм.
        row.cells[4].text = item.get('Step', '-')                   # Шаг
        row.cells[5].text = item.get('DefaultValue', '')            # По умолчанию
        row.cells[6].text = str(item.get('CurrentSetting', ''))     # Уставка

        # Жирный шрифт для ячейки "Уставка"
        cell_ust = row.cells[6]
        for paragraph in cell_ust.paragraphs:
            for run in paragraph.runs:
                run.bold = True

        # Выравнивание
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                paragraph.style = 'Текст таблицы'

        # !!! ИСПРАВЛЕНО: Color с большой буквы !!!
        is_changed = item.get('Color', '0')
        
        # Цветовая индикация
        if is_changed == '1' or is_changed == 1:
            shading_color = "FFC000"  # Желтый
        elif is_changed == '2' or is_changed == 2:
            shading_color = "FF0000"  # Красный
        else:
            shading_color = None

        if shading_color:
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shading_color}"/>')
            cell_ust._tc.get_or_add_tcPr().append(shading_elm)

    # Стиль таблицы
    table.style = 'Стиль3'
    table.allow_autofit = False
    table.autofit = False
    
    # Ширина столбцов
    widths = [1.75, 0.98, 1.4, 0.48, 0.47, 1.0, 0.6]
    for i, width in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = Inches(width)    
    
    return doc