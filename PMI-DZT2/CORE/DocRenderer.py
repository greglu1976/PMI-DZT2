from docx import Document
from docx.shared import Mm, Inches, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from CORE.tables import set_repeat_table_header, set_cell_vertical_alignment
import re


class DocRenderer:
    def __init__(self, template_path):
        self.doc = Document(template_path + '/template_apx.docx')

    def proceed(self, modes_list, func_for_header, doc=None):
        # Определяем, с каким документом работаем: переданным или внутренним
        target_doc = doc if doc is not None else self.doc
        
        # Добавляем заголовок раздела
        mode_name = func_for_header.split('_')[0]
        target_doc.add_heading(f'Параметры для проверки {mode_name}', level=1)

        for mode_index, mode in enumerate(modes_list):
            # ... (ваш код без изменений, но заменяем self.doc на target_doc) ...
            mode_number = mode["mode_file"].replace('.xlsx', '').split('_')[-1]
            target_doc.add_heading(f'Режим №{mode_number}', level=2)

            if mode_index == 0:
                for block_name, setting_block in mode['settings'].items():
                    # Если заголовки блоков нужны только в основном документе, оставьте как есть.
                    # Но обычно они тоже должны идти в target_doc
                    target_doc.add_heading(setting_block["BlockFullRusName"], level=3)
                    
                    for func_name, function in setting_block['functions'].items():
                        # Передаем target_doc внутрь, если _add_table_settings тоже использует self.doc
                        # Если _add_table_settings жестко привязан к self.doc, то вариант с local var сложнее.
                        # В вашем текущем коде _add_table_settings использует self.doc напрямую.
                        # Поэтому вариант с заменой self.doc (как у вас) проще, но требует осторожности.
                        self._add_table_settings(function, mode['mode_file'], is_first_mode=True)
            
            else:
                has_changes = self._has_mode_changes(mode)
                
                if not has_changes:
                    paragraph = target_doc.add_paragraph(style='ЮИ_Обычный')
                    run = paragraph.add_run("Уставки и параметры задействованных функций идентичны предыдущему режиму.")
                    run.font.italic = True
                    continue
                
                for block_name, setting_block in mode['settings'].items():
                    has_block_changes = self._has_block_changes(setting_block)
                    
                    if not has_block_changes:
                        continue
                    
                    for func_name, function in setting_block['functions'].items():
                        has_func_changes = self._has_func_changes(function)
                        if not has_func_changes:
                            continue
                        
                        self._add_table_settings(function, mode['mode_file'], is_first_mode=False)

        # Если мы использовали временный документ, возвращаем его.
        # Если использовали self.doc, тоже возвращаем его.
        return target_doc  



    def _has_mode_changes(self, mode):
        """Проверяет, есть ли изменения в режиме"""
        for block_name, setting_block in mode['settings'].items():
            if self._has_block_changes(setting_block):
                return True
        return False

    def _has_block_changes(self, setting_block):
        """Проверяет, есть ли изменения в блоке"""
        for func_name, function in setting_block['functions'].items():
            if self._has_func_changes(function):
                return True
        return False

    def _has_func_changes(self, function):
        """Проверяет, есть ли изменения в функции"""
        for param in function.get('parameters', []):
            if param.get('Color') == '1':
                return True
        return False

    def _parse_options_string(self, text):
        """
        Преобразует строку вида "1 - Управляющее напряжение, 2 - Вольтметровая блокировка"
        в словарь {"1": "Управляющее напряжение", "2": "Вольтметровая блокировка"}
        """
        if not text or not isinstance(text, str):
            return {} , ""
        
        result = {}
        # Разделяем по запятой (учитываем возможные пробелы)
        pairs = re.split(r'\s*,\s*', text.strip())
        
        for pair in pairs:
            # Разделяем по " - " или "-" с возможными пробелами
            match = re.match(r'^\s*(\d+)\s*-\s*(.+?)\s*$', pair)
            if match:
                key = match.group(1)
                value = match.group(2).strip()
                result[key] = value
        
        return result, " / ".join(result.values())


    def _format_values(self, step_val, default, min_val, max_val, current):
        """
        Форматирует значения по шагу: выравнивает количество знаков и меняет точку на запятую.
        """
        # Определяем количество знаков после запятой по шагу
        decimals = len(step_val.split('.')[1]) if '.' in step_val else 0
        
        # Форматируем каждое значение
        default = f"{float(default):.{decimals}f}".replace('.', ',')
        min_val = f"{float(min_val):.{decimals}f}".replace('.', ',')
        max_val = f"{float(max_val):.{decimals}f}".replace('.', ',')
        step = f"{float(step_val):.{decimals}f}".replace('.', ',')
        current = f"{float(current):.{decimals}f}".replace('.', ',')
        
        return step, default, min_val, max_val, current


    def _add_table_settings(self, function, mode_file, is_first_mode=True):
        """Добавляет таблицу настроек"""
        func_name = function.get("NodeNameRu", "")
        func_full_name = function.get("FullNodeNameRu", "")

        # Заголовок функции
        header_func = f'{func_name} ({func_full_name})'
        if func_name == 'Общие уставки':
            header_func = 'Общие уставки'
        
        self.doc.add_paragraph(header_func, style='ЮИ_Таблица_Название')
        
        # Создаем таблицу
        table = self.doc.add_table(rows=1, cols=7)
        header_row = table.rows[0]
        header_row.height = Mm(5)
        set_repeat_table_header(header_row)

        # Заголовки столбцов
        headers = ['Параметр', 'Обозначение ФСУ', 'Значение / Диапазон', 'Ед. изм.', 'Шаг', 'Значение по умолчанию', 'Выставленная уставка']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell.paragraphs[0].style = 'Текст таблицы'
            set_cell_vertical_alignment(cell, align="center")

        # Заполнение данными
        for item in function.get("parameters", []):
            # ✅ Для не-первого режима пропускаем параметры без изменений
            if not is_first_mode and item.get('Color') != '1':
                continue

            # подготовка значений для вставки в таблицу
            note = item.get('Note')
            default = item.get('DefaultValue')
            current = str(item.get('CurrentSetting'))
            step = item.get('Step')

            diap = ""
            if note != '':
                option_dict, diap = self._parse_options_string(note)
                default_setting = option_dict.get(default)
                current_setting = option_dict.get(current)
                step = '-'
                
            else:
                step_val, default_val, min_val, max_val, current_val = self._format_values(step, default, item.get('Min'), item.get('Max'), current)
                diap = f"{min_val} ... {max_val}"
                default_setting = default_val
                current_setting = current_val
                step = step_val


            row = table.add_row()
            row.cells[0].text = f"{item.get('FullDescription', '')} ({item.get('Description', '')}) "
            row.cells[1].text = item.get('AppliedDescription') if item.get('AppliedDescription')!="" else "-"
            row.cells[2].text = diap
            row.cells[3].text = item.get('Units') if item.get('Units')!="" else "-"
            row.cells[4].text = step
            row.cells[5].text = default_setting
            row.cells[6].text = current_setting

            # Жирный шрифт для ячейки "Уставка"
            cell_ust = row.cells[6]
            for paragraph in cell_ust.paragraphs:
                for run in paragraph.runs:
                    #run.bold = True
                    run.italic = True

            # Выравнивание
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    paragraph.style = 'Текст таблицы'

            # Цветовая индикация (только для изменённых)
            if item.get('Color') == '1':
                shading_color = "00FF7B" #"11502F" #"FFC000"  # Желтый
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shading_color}"/>')
                cell_ust._tc.get_or_add_tcPr().append(shading_elm)

        table.style = 'Стиль3'
        table.allow_autofit = False
        table.autofit = False
        
        # Ширина столбцов
        widths = [1.35, 0.98, 1.35, 0.48, 0.47, 1.0, 1.05]
        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(width)

    def save_docx(self, name='Output/Уставки режимов.docx'):
        self.doc.save(name)